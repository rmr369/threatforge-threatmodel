"""
Word threat model document, for audit / ARB / compliance evidence.

Optional dependency: python-docx. If it is not installed the CLI falls back to
Markdown and says so rather than failing the run.
"""

from __future__ import annotations

from datetime import datetime, timezone
from typing import Optional

from ..model import Element, Severity, ThreatModel

STRIDE_FULL = {
    "S": "Spoofing", "T": "Tampering", "R": "Repudiation",
    "I": "Information Disclosure", "D": "Denial of Service",
    "E": "Elevation of Privilege",
}

LEVEL_RGB = {
    "critical": (0xC0, 0x1B, 0x1B), "high": (0xC2, 0x41, 0x0C),
    "medium": (0x92, 0x6C, 0x0A), "low": (0x1D, 0x4E, 0xD8), "info": (0x47, 0x55, 0x69),
}


def available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def render(model: ThreatModel, path: str, *, max_findings: int = 50) -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    # -- title page --------------------------------------------------------
    doc.add_heading(f"Threat Model: {model.project}", 0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        f"Generated {datetime.now(timezone.utc).strftime('%d %B %Y')} · "
        f"ThreatForge · evidence-based STRIDE analysis of infrastructure-as-code")
    run.italic = True
    run.font.size = Pt(9.5)

    counts = model.counts()

    # -- executive summary -------------------------------------------------
    doc.add_heading("1. Executive summary", 1)
    exposed = sum(1 for a in model.assets.values() if a.facts.get("internet_reachable"))
    stores = len([a for a in model.assets.values()
                  if a.element == Element.DATA_STORE and a.sensitivity >= 4])
    doc.add_paragraph(
        f"This threat model covers {len(model.assets)} assets connected by "
        f"{len(model.flows)} data flows across {len(model.boundaries)} trust boundaries. "
        f"{exposed} assets are reachable from an untrusted network and {stores} hold "
        f"sensitive data. Analysis produced {len(model.active_findings)} findings, each "
        f"backed by a specific configuration observed in the scanned sources.")

    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = "Risk level"
    t.rows[0].cells[1].text = "Findings"
    for lvl in ("critical", "high", "medium", "low", "info"):
        if counts.get(lvl):
            row = t.add_row().cells
            row[0].text = lvl.title()
            row[1].text = str(counts[lvl])

    if model.attack_paths:
        top = model.attack_paths[0]
        tgt = model.assets.get(top.target)
        doc.add_paragraph(
            f"The highest-scoring attack path reaches "
            f"{tgt.display if tgt else top.target} in {len(top.hops) - 1} hops "
            f"(score {top.score}).", style="Intense Quote")

    # -- scope -------------------------------------------------------------
    doc.add_heading("2. Scope and limitations", 1)
    ing = model.metadata.get("ingestors", {})
    if ing:
        doc.add_paragraph("Sources analysed:")
        for name, stats in ing.items():
            if stats.get("assets"):
                doc.add_paragraph(
                    f"{name}: {stats.get('files', 0)} files, "
                    f"{stats.get('assets', 0)} assets", style="List Bullet")
    doc.add_paragraph(
        "Out of scope: application source code, container image contents and known "
        "vulnerabilities, runtime behaviour, admission control or service mesh policy "
        "applied outside the scanned sources, and resources created out-of-band. The "
        "absence of a finding is not evidence that a risk does not exist.")

    # -- trust boundaries --------------------------------------------------
    doc.add_heading("3. Trust boundaries", 1)
    tb = doc.add_table(rows=1, cols=4)
    tb.style = "Light Grid Accent 1"
    for i, h in enumerate(("Boundary", "Kind", "Trust", "Assets")):
        tb.rows[0].cells[i].text = h
    for b in sorted(model.boundaries.values(), key=lambda x: x.trust_level):
        c = tb.add_row().cells
        c[0].text, c[1].text = b.name, b.kind
        c[2].text, c[3].text = str(b.trust_level), str(len(b.members))

    # -- attack paths ------------------------------------------------------
    doc.add_heading("4. Attack paths", 1)
    if not model.attack_paths:
        doc.add_paragraph(
            "No complete path was found from an untrusted entry point to a crown-jewel "
            "asset in the modelled graph.")
    else:
        for i, ap in enumerate(model.attack_paths[:5], start=1):
            entry = model.assets.get(ap.entry)
            tgt = model.assets.get(ap.target)
            doc.add_heading(
                f"4.{i} {entry.display if entry else ap.entry} → "
                f"{tgt.display if tgt else ap.target}", 2)
            doc.add_paragraph(
                f"Score {ap.score} ({ap.level.value}), {len(ap.hops)} hops, "
                f"{len(ap.findings)} enabling findings.")
            for step in ap.narrative:
                doc.add_paragraph(step, style="List Number")

    # -- STRIDE ------------------------------------------------------------
    doc.add_heading("5. STRIDE coverage", 1)
    st = doc.add_table(rows=1, cols=3)
    st.style = "Light Grid Accent 1"
    for i, h in enumerate(("Category", "Findings", "Highest risk")):
        st.rows[0].cells[i].text = h
    for letter, name in STRIDE_FULL.items():
        fs = [f for f in model.active_findings if letter in f.stride]
        c = st.add_row().cells
        c[0].text = f"{letter} — {name}"
        c[1].text = str(len(fs))
        c[2].text = str(max((f.risk_score for f in fs), default=0))

    # -- findings register --------------------------------------------------
    doc.add_heading("6. Findings register", 1)
    shown = model.active_findings[:max_findings]
    if len(model.active_findings) > max_findings:
        doc.add_paragraph(
            f"Showing the {max_findings} highest-risk of "
            f"{len(model.active_findings)} findings. The full register is in "
            f"threat-model.json.").italic = True

    for f in shown:
        h = doc.add_heading(f"{f.risk_score}/25 — {f.title}", 2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(*LEVEL_RGB.get(f.risk_level.value,
                                                         (0x47, 0x55, 0x69)))
        meta = doc.add_paragraph()
        mr = meta.add_run(
            f"{f.rule_id} · {f.risk_level.value.upper()} · {f.component} · "
            f"confidence {f.confidence.value} · STRIDE "
            f"{', '.join(f.stride) or '—'}")
        mr.font.size = Pt(9)
        mr.italic = True

        if f.description:
            doc.add_paragraph(" ".join(f.description.split()))

        doc.add_paragraph("Evidence", style="Intense Quote")
        for e in f.evidence:
            loc = ""
            if e.source and e.source.file:
                loc = f" [{e.source.file}" + (f":{e.source.line}" if e.source.line else "") + "]"
            obs = f" (observed: {e.observed})" if e.observed not in (None, "") else ""
            doc.add_paragraph(f"{e.description}{obs}{loc}", style="List Bullet")

        r = f.risk
        exposure = ("not reachable from an external entity" if r.exposure_hops is None
                    else f"{r.exposure_hops} hop(s) from the internet")
        doc.add_paragraph(
            f"Risk rationale: likelihood {r.likelihood} × impact {r.impact} = "
            f"{f.risk_score}. Exposure: {exposure}. Blast radius: {r.blast_radius}. "
            f"Data sensitivity: {r.sensitivity}.")
        for n in r.notes:
            doc.add_paragraph(n, style="List Bullet")

        if f.remediation:
            doc.add_paragraph(
                f"Remediation: {f.remediation.summary} "
                f"(effort {f.remediation.effort}, breaking risk "
                f"{f.remediation.breaking_risk})")
            if f.remediation.guidance:
                doc.add_paragraph(" ".join(f.remediation.guidance.split()))
            if f.remediation.patch:
                cp = doc.add_paragraph()
                cr = cp.add_run(f.remediation.patch.rstrip())
                cr.font.name = "Consolas"
                cr.font.size = Pt(8.5)

        if f.references:
            refs = "; ".join(f"{k.upper()}: {', '.join(v)}"
                             for k, v in f.references.items() if v)
            if refs:
                rp = doc.add_paragraph(refs)
                rp.runs[0].font.size = Pt(8.5)
                rp.runs[0].italic = True

    # -- accepted risks -----------------------------------------------------
    supp = [f for f in model.findings if f.suppressed]
    if supp:
        doc.add_heading("7. Accepted risks and suppressions", 1)
        at = doc.add_table(rows=1, cols=3)
        at.style = "Light Grid Accent 1"
        for i, hh in enumerate(("Rule", "Component", "Reason")):
            at.rows[0].cells[i].text = hh
        for f in supp[:80]:
            c = at.add_row().cells
            c[0].text, c[1].text = f.rule_id, f.component
            c[2].text = f.suppression_reason or ""

    doc.save(path)
    return path
